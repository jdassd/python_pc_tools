import fitz
import os
import io
import tempfile

def pdf_to_images(pdf_path, output_folder):
    """
    Converts each page of a PDF to a PNG image.
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    doc = fitz.open(pdf_path)
    for page_index in range(len(doc)):
        page = doc[page_index]
        pix = page.get_pixmap()
        output_path = os.path.join(output_folder, f"page_{page_index + 1}.png")
        pix.save(output_path)
    doc.close()
    return f"Successfully converted {len(doc)} pages to images in {output_folder}"

def merge_pdfs(pdf_files, output_path):
    """
    Merges multiple PDF files into one.
    """
    result_pdf = fitz.open()
    for pdf_file in pdf_files:
        with fitz.open(pdf_file) as doc:
            result_pdf.insert_pdf(doc)
    result_pdf.save(output_path)
    result_pdf.close()
    return f"Successfully merged {len(pdf_files)} files into {output_path}"

def split_pdf(pdf_path, ranges, output_folder):
    """
    Splits a PDF into multiple files based on page ranges.
    Ranges should be a list of tuples, e.g., [(1, 5), (6, 10)]
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        
    original_doc = fitz.open(pdf_path)
    for i, (start, end) in enumerate(ranges):
        new_doc = fitz.open()
        new_doc.insert_pdf(original_doc, from_page=start-1, to_page=end-1)
        output_path = os.path.join(output_folder, f"split_{i+1}.pdf")
        new_doc.save(output_path)
        new_doc.close()
    original_doc.close()
    return f"Successfully split PDF into {len(ranges)} files in {output_folder}"

def _get_page_conversion_strategy(page):
    """
    Analyzes a page to decide the best conversion strategy.
    'layout' for pages with tables/vectors. 'image' for scanned pages.
    """
    if len(page.get_drawings()) > 5: # Heuristic for tables/forms
        return 'layout'
    
    text = page.get_text("text").strip()
    if len(text) < 100 and len(page.get_images()) > 0:
        return 'image'
        
    return 'layout'

def _convert_page_layout_aware(pdf_path, page_num):
    """
    Converts a single page using pdf2docx and returns the path to a temp docx file.
    This function is self-contained to prevent file handle leaks.
    """
    from pdf2docx import Converter
    import tempfile

    temp_docx_path = None
    try:
        # Create a temporary file for the DOCX output
        temp_docx_fd, temp_docx_path = tempfile.mkstemp(suffix=".docx")
        os.close(temp_docx_fd)

        # Convert the specific page range directly
        cv = Converter(pdf_path)
        cv.convert(temp_docx_path, start=page_num, end=page_num)
        cv.close()
        
        return temp_docx_path
    except Exception as e:
        # Ensure cleanup if conversion fails
        if temp_docx_path and os.path.exists(temp_docx_path):
            os.remove(temp_docx_path)
        raise e

def pdf_to_word(pdf_path, output_path):
    """
    Converts a PDF to Word using a robust, page-by-page strategy that correctly
    handles file resources to prevent PermissionError on Windows.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT, WD_SECTION
        from docx.shared import Inches
    except ImportError as e:
        return f"Error: A required library is not installed. {e}"

    # 检查PDF文件是否存在
    if not os.path.exists(pdf_path):
        return f"Error: PDF 文件不存在: {pdf_path}"

    main_pdf_doc = None
    try:
        main_pdf_doc = fitz.open(pdf_path)
        word_doc = Document()

        for page_num, page in enumerate(main_pdf_doc):
            if page_num > 0:
                word_doc.add_section(WD_SECTION.NEW_PAGE)
            section = word_doc.sections[-1]

            # Set page orientation and size to match the PDF
            width_inch = page.rect.width / 72
            height_inch = page.rect.height / 72
            if width_inch > height_inch:
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(height_inch)
                section.page_height = Inches(width_inch)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(width_inch)
                section.page_height = Inches(height_inch)

            strategy = _get_page_conversion_strategy(page)

            if strategy == 'layout':
                temp_docx_path = None
                try:
                    temp_docx_path = _convert_page_layout_aware(pdf_path, page_num)
                    
                    # Read the temp docx into a memory stream to avoid file locks
                    with open(temp_docx_path, 'rb') as f:
                        docx_stream = io.BytesIO(f.read())
                    
                    # Open the document from the memory stream
                    temp_word_doc = Document(docx_stream)
                    for element in temp_word_doc.element.body:
                        if not element.tag.endswith('sectPr'):
                            word_doc.element.body.append(element)
                finally:
                    # Immediate cleanup
                    if temp_docx_path and os.path.exists(temp_docx_path):
                        os.remove(temp_docx_path)
            
            elif strategy == 'image':
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                pix = page.get_pixmap(dpi=200)
                img_stream = io.BytesIO(pix.tobytes("png"))
                available_width = section.page_width - section.left_margin - section.right_margin
                word_doc.add_picture(img_stream, width=available_width)

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        word_doc.save(output_path)
        return f"Successfully converted {pdf_path} to {output_path} with advanced page analysis."

    except Exception as e:
        return f"An error occurred during conversion: {e}"
    finally:
        if main_pdf_doc:
            main_pdf_doc.close()