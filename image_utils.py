from PIL import Image, ImageDraw, ImageFont, ImageEnhance, ImageOps, ImageFilter
from PIL.ExifTags import TAGS
from docx import Document
from docx.shared import Inches
import os
import time

def convert_image_format(input_path, output_path, format):
    """Converts an image to a different format."""
    try:
        img = Image.open(input_path)
        img.save(output_path, format=format)
        return f"Successfully converted {input_path} to {output_path}"
    except Exception as e:
        return f"Error converting image: {e}"

def resize_image(input_path, output_path, size):
    """Resizes an image."""
    try:
        img = Image.open(input_path)
        resized_img = img.resize(size)
        resized_img.save(output_path)
        return f"Successfully resized {input_path} to {size}"
    except Exception as e:
        return f"Error resizing image: {e}"

def crop_image(input_path, output_path, box):
    """Crops an image."""
    try:
        img = Image.open(input_path)
        cropped_img = img.crop(box)
        cropped_img.save(output_path)
        return f"Successfully cropped {input_path}"
    except Exception as e:
        return f"Error cropping image: {e}"

def add_watermark(input_path, output_path, text, position, font_path=None, font_size=20, color=(255, 255, 255, 128)):
    """Adds a text watermark to an image."""
    try:
        img = Image.open(input_path).convert("RGBA")
        txt = Image.new("RGBA", img.size, (255, 255, 255, 0))

        try:
            font = ImageFont.truetype(font_path, font_size)
        except IOError:
            font = ImageFont.load_default()

        draw = ImageDraw.Draw(txt)
        draw.text(position, text, font=font, fill=color)

        watermarked_img = Image.alpha_composite(img, txt)
        watermarked_img.save(output_path)
        return f"Successfully added watermark to {input_path}"
    except Exception as e:
        return f"Error adding watermark: {e}"

# 新增功能

def get_image_exif(image_path):
    """获取图片EXIF信息"""
    try:
        img = Image.open(image_path)
        exifdata = img.getexif()
        
        if not exifdata:
            return {"message": "该图片没有EXIF信息"}, None
        
        exif_dict = {}
        for tag_id, value in exifdata.items():
            tag = TAGS.get(tag_id, tag_id)
            exif_dict[tag] = str(value)
        
        # 添加基本图片信息
        basic_info = {
            "文件名": os.path.basename(image_path),
            "图片格式": img.format,
            "图片模式": img.mode,
            "图片尺寸": f"{img.width} x {img.height}",
            "文件大小": f"{os.path.getsize(image_path)} bytes"
        }
        
        return {"basic_info": basic_info, "exif_data": exif_dict}, None
        
    except Exception as e:
        return None, str(e)

def batch_rename_images(directory, name_pattern="IMG", use_creation_time=False):
    """批量重命名图片文件"""
    try:
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
        image_files = []
        
        # 获取所有图片文件
        for filename in os.listdir(directory):
            if any(filename.lower().endswith(ext) for ext in image_extensions):
                file_path = os.path.join(directory, filename)
                if use_creation_time:
                    # 使用创建时间排序
                    creation_time = os.path.getctime(file_path)
                    image_files.append((filename, file_path, creation_time))
                else:
                    image_files.append((filename, file_path, filename.lower()))
        
        # 排序
        if use_creation_time:
            image_files.sort(key=lambda x: x[2])  # 按创建时间排序
        else:
            image_files.sort(key=lambda x: x[2])  # 按文件名排序
        
        renamed_files = []
        errors = []
        
        for i, (old_name, file_path, _) in enumerate(image_files, 1):
            _, ext = os.path.splitext(old_name)
            new_name = f"{name_pattern}_{i:04d}{ext}"
            new_path = os.path.join(directory, new_name)
            
            if os.path.exists(new_path):
                errors.append(f"{old_name}: 目标文件已存在")
                continue
            
            try:
                os.rename(file_path, new_path)
                renamed_files.append((old_name, new_name))
            except Exception as e:
                errors.append(f"{old_name}: {str(e)}")
        
        return {"renamed": renamed_files, "errors": errors}, None
        
    except Exception as e:
        return None, str(e)

def rotate_flip_image(input_path, output_path, operation):
    """旋转或翻转图片"""
    try:
        img = Image.open(input_path)
        
        if operation == "rotate_90":
            result_img = img.rotate(90, expand=True)
        elif operation == "rotate_180":
            result_img = img.rotate(180)
        elif operation == "rotate_270":
            result_img = img.rotate(270, expand=True)
        elif operation == "flip_horizontal":
            result_img = img.transpose(Image.Transpose.FLIP_LEFT_RIGHT)
        elif operation == "flip_vertical":
            result_img = img.transpose(Image.Transpose.FLIP_TOP_BOTTOM)
        else:
            return None, f"不支持的操作: {operation}"
        
        result_img.save(output_path)
        return f"成功{operation}图片: {input_path}", None
        
    except Exception as e:
        return None, str(e)

def concat_images(image_paths, output_path, direction="horizontal", spacing=0):
    """拼接图片"""
    try:
        images = [Image.open(path) for path in image_paths]
        
        if direction == "horizontal":
            # 水平拼接
            total_width = sum(img.width for img in images) + spacing * (len(images) - 1)
            max_height = max(img.height for img in images)
            
            result = Image.new('RGB', (total_width, max_height), (255, 255, 255))
            
            x_offset = 0
            for img in images:
                result.paste(img, (x_offset, 0))
                x_offset += img.width + spacing
                
        else:  # vertical
            # 垂直拼接
            max_width = max(img.width for img in images)
            total_height = sum(img.height for img in images) + spacing * (len(images) - 1)
            
            result = Image.new('RGB', (max_width, total_height), (255, 255, 255))
            
            y_offset = 0
            for img in images:
                result.paste(img, (0, y_offset))
                y_offset += img.height + spacing
        
        result.save(output_path)
        return f"成功拼接{len(images)}张图片", None
        
    except Exception as e:
        return None, str(e)

def adjust_image_properties(input_path, output_path, brightness=1.0, contrast=1.0, saturation=1.0, sharpness=1.0):
    """调整图片属性"""
    try:
        img = Image.open(input_path)
        
        # 调整亮度
        if brightness != 1.0:
            enhancer = ImageEnhance.Brightness(img)
            img = enhancer.enhance(brightness)
        
        # 调整对比度
        if contrast != 1.0:
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(contrast)
        
        # 调整饱和度
        if saturation != 1.0:
            enhancer = ImageEnhance.Color(img)
            img = enhancer.enhance(saturation)
        
        # 调整锐度
        if sharpness != 1.0:
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(sharpness)
        
        img.save(output_path)
        return f"成功调整图片属性: {input_path}", None
        
    except Exception as e:
        return None, str(e)

def apply_image_filter(input_path, output_path, filter_type):
    """应用图片滤镜"""
    try:
        img = Image.open(input_path)
        
        if filter_type == "blur":
            result_img = img.filter(ImageFilter.BLUR)
        elif filter_type == "detail":
            result_img = img.filter(ImageFilter.DETAIL)
        elif filter_type == "edge_enhance":
            result_img = img.filter(ImageFilter.EDGE_ENHANCE)
        elif filter_type == "emboss":
            result_img = img.filter(ImageFilter.EMBOSS)
        elif filter_type == "find_edges":
            result_img = img.filter(ImageFilter.FIND_EDGES)
        elif filter_type == "smooth":
            result_img = img.filter(ImageFilter.SMOOTH)
        elif filter_type == "sharpen":
            result_img = img.filter(ImageFilter.SHARPEN)
        elif filter_type == "grayscale":
            result_img = img.convert('L')
        elif filter_type == "invert":
            result_img = ImageOps.invert(img)
        else:
            return None, f"不支持的滤镜类型: {filter_type}"
        
        result_img.save(output_path)
        return f"成功应用{filter_type}滤镜", None
        
    except Exception as e:
        return None, str(e)

def batch_process_images(directory, operation, **kwargs):
    """批量处理图片"""
    try:
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp']
        processed_files = []
        errors = []
        
        for filename in os.listdir(directory):
            if any(filename.lower().endswith(ext) for ext in image_extensions):
                input_path = os.path.join(directory, filename)
                name, ext = os.path.splitext(filename)
                output_filename = f"{name}_{operation}{ext}"
                output_path = os.path.join(directory, output_filename)
                
                try:
                    if operation == "compress":
                        # 压缩图片
                        img = Image.open(input_path)
                        quality = kwargs.get('quality', 85)
                        img.save(output_path, optimize=True, quality=quality)
                        processed_files.append((filename, output_filename))
                        
                    elif operation == "resize":
                        # 批量调整大小
                        size = kwargs.get('size', (800, 600))
                        result, error = resize_image(input_path, output_path, size)
                        if error:
                            errors.append(f"{filename}: {error}")
                        else:
                            processed_files.append((filename, output_filename))
                            
                    elif operation in ["rotate_90", "rotate_180", "rotate_270", "flip_horizontal", "flip_vertical"]:
                        # 批量旋转翻转
                        result, error = rotate_flip_image(input_path, output_path, operation)
                        if error:
                            errors.append(f"{filename}: {error}")
                        else:
                            processed_files.append((filename, output_filename))
                            
                    elif operation in ["blur", "detail", "edge_enhance", "emboss", "find_edges", "smooth", "sharpen", "grayscale", "invert"]:
                        # 批量应用滤镜
                        result, error = apply_image_filter(input_path, output_path, operation)
                        if error:
                            errors.append(f"{filename}: {error}")
                        else:
                            processed_files.append((filename, output_filename))
                    
                except Exception as e:
                    errors.append(f"{filename}: {str(e)}")
        
        return {"processed": processed_files, "errors": errors}, None
        
    except Exception as e:
        return None, str(e)

def images_to_word(image_paths, output_path, existing_word_file=None, start_page=1):
    """将图片插入Word文档"""
    try:
        if existing_word_file and os.path.exists(existing_word_file):
            doc = Document(existing_word_file)
            is_new_file = False
        else:
            doc = Document()
            is_new_file = True
        
        # 如果需要从指定页开始插入且不是新文件
        if not is_new_file and start_page > 1:
            # 计算现有页数
            existing_pages = len(doc.paragraphs) // 10 + 1  # 粗略估算
            if existing_pages < start_page:
                # 添加分页符直到到达指定页
                for _ in range(start_page - existing_pages):
                    doc.add_page_break()
            else:
                # 在现有内容后添加分页符
                doc.add_page_break()
        
        for i, image_path in enumerate(image_paths):
            if not os.path.exists(image_path):
                continue
            
            try:
                # 获取图片尺寸并保持比例
                with Image.open(image_path) as img:
                    img_width, img_height = img.size
                    aspect_ratio = img_width / img_height
                    
                    # 设置最大尺寸限制（Word页面可用空间）
                    max_width = 6.0  # inches
                    max_height = 8.0  # inches
                    
                    # 根据比例计算合适的尺寸
                    if aspect_ratio > max_width / max_height:
                        # 图片较宽，以宽度为准
                        width_inch = min(max_width, img_width / 100.0)
                        height_inch = width_inch / aspect_ratio
                    else:
                        # 图片较高，以高度为准
                        height_inch = min(max_height, img_height / 100.0)
                        width_inch = height_inch * aspect_ratio
                
                # 如果不是第一张图片，添加分页符
                if i > 0:
                    doc.add_page_break()
                
                # 添加图片到新的段落
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(width_inch), height=Inches(height_inch))
                
            except Exception as img_error:
                # 如果图片插入失败，添加错误信息
                if i > 0:
                    doc.add_page_break()
                doc.add_paragraph(f"无法插入图片: {os.path.basename(image_path)} - {str(img_error)}")
        
        doc.save(output_path)
        
        if existing_word_file and output_path == existing_word_file:
            return f"成功将{len(image_paths)}张图片插入到现有Word文档，从第{start_page}页开始"
        else:
            return f"成功创建Word文档，包含{len(image_paths)}张图片"
        
    except Exception as e:
        return f"创建Word文档时出错: {str(e)}"